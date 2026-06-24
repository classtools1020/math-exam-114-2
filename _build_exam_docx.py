# -*- coding: utf-8 -*-
"""樂學教室 6/25 期末考 — Word(.docx) 版生成器
範圍：翰林版 3-5～4-3｜每題一圖（向量 PNG）｜SEL + 生活連結 + 素養
產出：學生卷.docx / 教師答案卷.docx"""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).parent
OUT = ROOT / "樂學教室_6-25期末考_正式版"
FIG = Path(r"C:/Users/celdi/AppData/Local/Temp/lx_exam/figpng")
FONT = "Microsoft JhengHei"

INK = "243049"; INKSOFT = "5A6B86"; TEAL = "0E8A8A"; CORAL = "C0392B"
GOLD = "B07A0C"; VIOLET = "6C5CE0"; SKY = "2B7FC9"; WHITE = "FFFFFF"
FAINT = "F4F7FB"; TEALF = "E2F4F3"; CORALF = "FDEBE4"; GOLDF = "FBF0D6"
VIOF = "EEEBFD"; SKYF = "E6F1FB"; NAVYF = "EAF1F8"

# ---- 題目資料（對齊原班翰林版 3-5～4-3）----
Q = [
 dict(n=1, t="三角形組得成嗎？", pts=8, fig="tri", lv="A", life="動漫房 RGB 三段燈框",
   sel="先看圖找三邊，再決定第三邊能多長。你做得到。",
   prompt=["三段燈條長度是 4、9、2x − 1。", "要圍成一個三角形，求 x 的範圍。"],
   steps=["① 第三邊要比「兩邊差」大、比「兩邊和」小：9 − 4 ＜ 2x − 1 ＜ 9 + 4",
     "② ____ ＜ 2x − 1 ＜ ____", "③ 三段同時 +1、再 ÷2 → ____ ＜ x ＜ ____"],
   hint="兩邊差 ＜ 第三邊 ＜ 兩邊和。", ans="3 ＜ x ＜ 7",
   sol="5 ＜ 2x−1 ＜ 13 → 6 ＜ 2x ＜ 14 → 3 ＜ x ＜ 7"),
 dict(n=2, t="拼出最大的角", pts=8, fig="angle", lv="A", life="動漫社全等三角形徽章",
   sel="角度也能拼圖。先補滿三角形，再想拼起來的樣子。",
   prompt=["一種三角形紙片，兩個內角是 38° 與 72°。", "用兩塊這種紙片拼成平行四邊形，最大的內角是幾度？"],
   steps=["① 第三個角 = 180° − 38° − 72° = ____°", "② 平行四邊形最大內角 = 180° − 最小角(38°) = ____°"],
   hint="三角形最小角的補角，就是平行四邊形最大的角。", ans="142°",
   sol="第三角 70°；最小角 38°，180−38 = 142°"),
 dict(n=3, t="斜坡的轉角（內錯角）", pts=8, fig="alt", lv="A", life="滑板場兩道平行護欄",
   sel="兩條平行線是好朋友，會把角度一模一樣傳過去。",
   prompt=["兩道護欄 L₁ ∥ L₂，一條斜桿穿過。", "圖中 58° 與 x 是內錯角，求 x。"],
   steps=["內錯角相等 → x = ____°"], hint="x 在兩平行線中間、兩側 → 內錯角 → 相等。",
   ans="58°", sol="內錯角相等，x = 58°"),
 dict(n=4, t="護欄的同側角", pts=8, fig="same", lv="A", life="滑板場護欄轉折處",
   sel="同一邊的兩個角會湊成 180°，像合作的隊友。",
   prompt=["L₁ ∥ L₂，63° 與 x 在截線同一側、兩線之間。", "求 x。"],
   steps=["同側內角互補 → 63° + x = 180°", "x = ____°"], hint="同側內角 → 兩角相加 = 180°。",
   ans="117°", sol="180 − 63 = 117°"),
 dict(n=5, t="舞台板面積放大術", pts=9, fig="paArea", lv="B", life="音樂祭 LED 舞台板",
   sel="小三角形藏著大面板的祕密，跟著比例放大就好。",
   prompt=["平行四邊形 ABCD，E、F 是 AB、BC 的中點。", "若 △BEF 面積 = 3，求平行四邊形 ABCD 面積。"],
   steps=["① △BEF 用了「半個底 × 半個高」 → 是整個平行四邊形的 1/8",
     "② 平行四邊形面積 = 3 × ____ = ____"],
   hint="中點各砍一半：1/2 × 1/2，三角形再 ÷2，合起來是 1/8。", ans="24",
   sol="△BEF = 1/8 × □ → □ = 3 × 8 = 24",
   plus="再想一步：整個平行四邊形可以切成幾個和 △BEF 一樣大的小三角形？", plusans="8 個"),
 dict(n=6, t="對角線偵探（打 ○／×）", pts=9, fig="family", lv="B", life="手機相框與收藏卡版型",
   sel="每個四邊形都有性格，對角線就是它的指紋。",
   prompt=["看圖判斷，每句對的打 ○、錯的打 ×。"],
   judge=["(1) 平行四邊形的對角線互相平分。（　　）",
     "(2) 長方形的對角線等長，而且互相垂直。（　　）",
     "(3) 菱形的對角線互相垂直平分。（　　）",
     "(4) 正方形同時具有長方形與菱形的對角線性質。（　　）"],
   hint="長方形是「等長 + 互相平分」，不一定垂直喔。", ans="○、×、○、○",
   sol="(2) 長方形對角線不垂直，故為 ×"),
 dict(n=7, t="做一只會飛的風箏", pts=9, fig="kite", lv="B", life="校慶風箏工坊",
   sel="把風箏切成上下兩個直角三角形，畢氏定理就來幫忙。",
   prompt=["箏形 ABCD，AC 垂直平分 BD。AB = 10、BC = 17、BD = 16。", "求對角線 AC 與風箏面積。"],
   steps=["① BO = 16 ÷ 2 = ____", "② AO = √(10² − 8²) = ____　CO = √(17² − 8²) = ____",
     "③ AC = AO + CO = ____", "④ 面積 = AC × BD ÷ 2 = ____"],
   hint="對角線互相垂直 → 面積 = 兩對角線相乘 ÷ 2。", ans="AC = 21；面積 = 168",
   sol="BO=8, AO=6, CO=15, AC=21, 面積=21×16÷2=168",
   plus="再想一步：這只風箏的周長是多少？", plusans="2×(10+17) = 54"),
 dict(n=8, t="中間那層架要多長？", pts=8, fig="trapMid", lv="B", life="手搖飲店梯形展示架",
   sel="中間層板 = 上下兩底的平均，公平地剛剛好。",
   prompt=["梯形展示架，上底 DC = 6、下底 AB = 14。", "E、F 是兩腰中點，求中層板 EF。"],
   steps=["中點連線 = (上底 + 下底) ÷ 2", "EF = (6 + 14) ÷ 2 = ____"],
   hint="兩底相加再除以 2。", ans="10", sol="(6+14)÷2 = 10"),
 dict(n=9, t="螢幕面積一次算出", pts=8, fig="trapArea", lv="B", life="音樂祭梯形主螢幕",
   sel="學會中點線，面積就是「中點線 × 高」，超省力。",
   prompt=["梯形主螢幕，上底 7、下底 11、高 8。", "先求中點線 EF，再求面積。"],
   steps=["① EF = (7 + 11) ÷ 2 = ____", "② 面積 = EF × 高 = ____ × 8 = ____"],
   hint="梯形面積 = 中點線 × 高。", ans="EF = 9；面積 = 72", sol="(7+11)÷2=9；9×8=72",
   plus="再想一步：如果高變成 2 倍，面積會變成幾倍？", plusans="2 倍"),
 dict(n=10, t="電競平台鋪多大？", pts=9, fig="isos", lv="C", star=True, life="電競舞台等腰梯形平台",
   sel="切下兩個全等直角三角形，高度就藏在裡面。挑戰看看！",
   prompt=["等腰梯形 ABCD，AB = 25、CD = 7、AD = BC = 15。", "求這座平台的面積。"],
   steps=["① 左右各凸出 (25 − 7) ÷ 2 = ____", "② 高 = √(15² − 9²) = ____",
     "③ 面積 = (25 + 7) × 高 ÷ 2 = ____"],
   hint="先算左右凸出 9，再用畢氏求高 12。", ans="192",
   sol="凸出9，高=√(225−81)=12，(25+7)×12÷2=192"),
 dict(n=11, t="下底被藏起來了", pts=8, fig="trapRev", lv="C", star=True, life="夾娃娃機梯形展示層",
   sel="公式可以倒著用，把未知數請出來。",
   prompt=["梯形中點連線 EF = 12，上底 DC = 7。", "反推下底 AB。"],
   steps=["12 = (7 + AB) ÷ 2", "兩邊 ×2 → 24 = 7 + AB", "AB = ____"],
   hint="先把 ÷2 還原成 ×2。", ans="17", sol="24 = 7 + AB → AB = 17"),
 dict(n=12, t="哪句話不能只信一半？", pts=8, fig="judge", lv="C", star=True, life="社群貼文卡片版型",
   sel="敢說『我覺得這裡怪怪的』，是數學最帥的勇氣。",
   prompt=["小明說：『只要兩條對角線等長，四邊形就是長方形。』", "請判斷對錯，並補上還缺的條件。"],
   steps=["判斷：□ 正確　　□ 錯誤", "反例：圖左的「____________」對角線也等長，卻不是長方形",
     "補上條件：對角線還要「________________」才行"],
   hint="等腰梯形對角線也等長；長方形還要『互相平分』。", ans="錯誤；還要互相平分",
   sol="等腰梯形是反例；長方形=對角線等長且互相平分"),
]
LVMAP = {"A": ("關卡 A", TEALF, TEAL), "B": ("關卡 B", SKYF, SKY), "C": ("挑戰關", GOLDF, GOLD)}

# ---------- helpers ----------
def set_run(r, size=11, bold=False, color=INK, italic=False):
    r.font.name = FONT
    rp = r._element.get_or_add_rPr()
    rf = rp.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rp.append(rf)
    for k in ("eastAsia", "ascii", "hAnsi"):
        rf.set(qn(f"w:{k}"), FONT)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)

def fmt(p, before=0, after=4, line=1.16, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    pf.keep_together = keep; pf.keep_with_next = keep

def add_p(doc_or_cell, s="", size=11, bold=False, color=INK, align=None,
          before=0, after=4, keep=False, italic=False):
    p = doc_or_cell.add_paragraph()
    fmt(p, before, after, keep=keep)
    if align is not None:
        p.alignment = align
    if s:
        set_run(p.add_run(s), size, bold, color, italic)
    return p

def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcpr.append(shd)
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)

def cell_margin(cell, top=70, bottom=70, start=110, end=110):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar"); tcpr.append(mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}"); mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'none'); b.append(e)
    tblPr.append(b)

def cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))

def shade_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def box(doc, fill, pad=(80,80,130,130)):
    """單格底色方塊，回傳 cell"""
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; t.columns[0].width = Cm(18.4)
    c = t.cell(0,0); shade(c, fill); cell_margin(c, *pad)
    return t, c

def band(doc, title, sub, fill=INK, tcolor=WHITE):
    t, c = box(doc, fill, (90,90,150,150))
    p = c.paragraphs[0]; fmt(p, before=2, after=0, keep=True)
    set_run(p.add_run(title), 14, True, tcolor)
    set_run(p.add_run("　" + sub), 10, False, "D7E0EC" if fill == INK else INKSOFT)

# ---------- 文件骨架 ----------
def setup(title, sub, teacher=False):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(1.2)
    sec.left_margin = sec.right_margin = Cm(1.3)
    sec.header_distance = sec.footer_distance = Cm(.5)
    st = doc.styles["Normal"]; st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(11); st.paragraph_format.space_after = Pt(4); st.paragraph_format.line_spacing = 1.16
    f = sec.footer.paragraphs[0]; f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(f.add_run("樂學教室期末考 ✦ 範圍對齊原班翰林版 3-5～4-3"), 8, False, INKSOFT)
    # 抬頭
    t, c = box(doc, INK, (160,160,170,170))
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p, after=3)
    set_run(p.add_run(title), 21, True, WHITE)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p2, after=2)
    set_run(p2.add_run(sub), 11.5, True, "BFD3EA")
    p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p3, after=0)
    set_run(p3.add_run("📅 6/25（四）期末考　🧭 12 關・滿分 100　🖼️ 每題都有圖　🌿 慢慢來，一步一步"), 9.5, False, "D7E0EC")
    add_p(doc, "", after=2)
    if teacher:
        t, c = box(doc, VIOF, (90,90,130,130))
        p = c.paragraphs[0]; fmt(p, after=0)
        set_run(p.add_run("教師版 ✦ 答案與評分　"), 11, True, VIOLET)
        set_run(p.add_run("每題附最短解法；性質正確、計算小錯，保留歷程分。對應原班段考核心概念，學生回原班可直接銜接。"), 9.8)
    else:
        info = doc.add_table(rows=1, cols=4); info.style = "Table Grid"
        for cc, s in zip(info.rows[0].cells, ("班級：______", "座號：______", "姓名：__________", "得分：______")):
            cell_margin(cc, 90, 90, 110, 110); cc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            pp = cc.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(pp.add_run(s), 10.5, True, INKSOFT)
        add_p(doc, "", after=2)
        t, c = box(doc, VIOF, (95,95,135,135))
        p = c.paragraphs[0]; fmt(p, after=0)
        set_run(p.add_run("開始前　"), 11, True, VIOLET)
        set_run(p.add_run("每一題都有圖。先看圖、圈出已知，再動筆。算錯不會扣光，過程對就有分——卡住很正常，舉手問也可以。"), 9.8)
    return doc

def work_box(doc, lines=4):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"; t.autofit = False; t.columns[0].width = Cm(18.4)
    c = t.cell(0, 0); cell_margin(c, 50, 60, 120, 120)
    p = c.paragraphs[0]; fmt(p, after=4)
    set_run(p.add_run("計算區（寫下你的算式，過程對就有分）"), 8.5, False, "A9B6C8")
    for _ in range(lines):
        lp = c.add_paragraph(); fmt(lp, after=8)
        set_run(lp.add_run("　"), 11)
    pa = c.add_paragraph(); fmt(pa, before=2, after=0)
    set_run(pa.add_run("答："), 10, True, INK)
    set_run(pa.add_run("　________________________"), 10, False, "A9B6C8")

def question(doc, q, teacher, variant):
    basic = (variant == "basic")
    lvname, lvfill, lvcol = LVMAP[q["lv"]]
    pts = 10 if basic else q["pts"]
    # 標題列
    p = add_p(doc, "", after=1, keep=True, before=7)
    set_run(p.add_run(f"{q['n']}. "), 13.5, True, lvcol)
    set_run(p.add_run(q["t"]), 13.5, True, INK)
    set_run(p.add_run(f"　（{pts} 分）"), 10, False, INKSOFT)
    if q.get("star"):
        set_run(p.add_run("　★ BOSS"), 10, True, GOLD)
    set_run(p.add_run(f"　〔{lvname}〕"), 9, True, lvcol)
    set_run(p.add_run(f"　{q['life']}"), 9, True, INKSOFT)
    # SEL（底色段落）— 兩版都保留，挑戰版語氣更精簡
    p = add_p(doc, "", after=2, keep=True); shade_para(p, FAINT)
    p.paragraph_format.left_indent = Cm(.15); p.paragraph_format.right_indent = Cm(.15)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(q["sel"]), 9.4, False, "3F5170")
    # 圖 + 內容（兩欄，整列不跨頁）
    tbl = doc.add_table(rows=1, cols=2); tbl.autofit = False; no_borders(tbl)
    tbl.columns[0].width = Cm(6.6); tbl.columns[1].width = Cm(11.8)
    cant_split(tbl.rows[0])
    lc, rc = tbl.rows[0].cells
    lc.width = Cm(6.6); rc.width = Cm(11.8)
    cell_margin(lc, 20, 20, 30, 50); cell_margin(rc, 20, 20, 70, 30)
    lc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    ip = lc.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(ip, after=0)
    ip.add_run().add_picture(str(FIG / f"{q['fig']}.png"), width=Cm(6.4 if basic else 6.2))
    # 右欄：題幹
    first = True
    for line in q["prompt"]:
        pp = rc.paragraphs[0] if first else rc.add_paragraph()
        fmt(pp, after=2, keep=True); set_run(pp.add_run(line), 11, False, INK); first = False
    # 挑戰版的加深小題
    if not basic and q.get("plus"):
        pp = rc.add_paragraph(); fmt(pp, after=2, keep=True)
        set_run(pp.add_run("★ " + q["plus"]), 10.5, True, GOLD)
    # 基礎版：步驟 / 判斷 + 完整提示
    if basic:
        for s in q.get("steps", q.get("judge", [])):
            pp = rc.add_paragraph(); fmt(pp, after=2, keep=True)
            pp.paragraph_format.left_indent = Cm(.2)
            set_run(pp.add_run(s), 10.5, False, INK)
        pp = rc.add_paragraph(); fmt(pp, before=2, after=1, keep=True)
        set_run(pp.add_run("💡 提示　"), 9.5, True, GOLD)
        set_run(pp.add_run(q["hint"]), 9.5, False, "6B5418")
    else:
        # 挑戰版：判斷題仍需列出選項；其餘只給圖與題目
        for s in q.get("judge", []):
            pp = rc.add_paragraph(); fmt(pp, after=2, keep=True)
            pp.paragraph_format.left_indent = Cm(.2)
            set_run(pp.add_run(s), 10.5, False, INK)
        if q.get("star"):
            pp = rc.add_paragraph(); fmt(pp, before=2, after=1, keep=True)
            set_run(pp.add_run("💡 一句線索　"), 9, True, GOLD)
            set_run(pp.add_run(q["hint"]), 9, False, "6B5418")
    # 答案 / 作答區
    if teacher:
        pp = rc.add_paragraph(); fmt(pp, before=2, after=0, keep=True)
        set_run(pp.add_run("✅ 答案　"), 10.5, True, TEAL)
        ans = q["ans"] + (("　｜★ " + q["plusans"]) if (not basic and q.get("plusans")) else "")
        set_run(pp.add_run(ans), 11, True, INK)
        pp2 = rc.add_paragraph(); fmt(pp2, after=0, keep=True)
        set_run(pp2.add_run("　解法：" + q["sol"]), 9, False, INKSOFT)
    elif basic:
        pp = rc.add_paragraph(); fmt(pp, before=3, after=0)
        set_run(pp.add_run("答：________________"), 10, False, "A9B6C8")
    # 挑戰版：學生卷在表格外補上完整計算區
    if not basic and not teacher:
        work_box(doc, lines=4)

def closing(doc):
    add_p(doc, "", after=2)
    t, c = box(doc, TEALF, (95,95,140,140))
    p = c.paragraphs[0]; fmt(p, after=2)
    set_run(p.add_run("交卷前，自己檢查三件事　"), 11, True, TEAL)
    set_run(p.add_run("（勾起來）"), 9.8, False, INKSOFT)
    p2 = c.add_paragraph(); fmt(p2, after=2)
    set_run(p2.add_run("☐ 內錯角／同側內角沒有搞混　　☐ 長方形 vs 菱形對角線分清楚　　☐ 面積有寫單位"), 10)
    p3 = c.add_paragraph(); fmt(p3, after=0)
    set_run(p3.add_run("這套關卡你練過了。下週回原班，遇到同樣的圖，照這個方法做就對了。"), 9.8, True, "3F5170")

def build(variant, teacher):
    basic = (variant == "basic")
    if basic:
        title = "樂學教室 ✦ 數學期末・基礎闖關"
        sub = "八年級第三次段考　範圍：翰林版 3-5～4-3　基礎版（步驟＋提示）"
        qs = [x for x in Q if x["n"] <= 10]           # 10 題 × 10 分
    else:
        title = "樂學教室 ✦ 數學期末・挑戰闖關"
        sub = "八年級第三次段考　範圍：翰林版 3-5～4-3　挑戰版（自己寫過程）"
        qs = list(Q)                                   # 12 題
    doc = setup(title, sub, teacher)
    band(doc, "關卡 A ✦ 暖身基礎", "三角形與平行線")
    for q in [x for x in qs if x["lv"] == "A"]:
        question(doc, q, teacher, variant)
    add_p(doc, "", after=2)
    band(doc, "關卡 B ✦ 四邊形與面積", "平行四邊形・箏形・梯形")
    for q in [x for x in qs if x["lv"] == "B"]:
        question(doc, q, teacher, variant)
    if any(x["lv"] == "C" for x in qs):
        add_p(doc, "", after=2)
        band(doc, "挑戰關 ✦ 銜接原班", "畢氏・逆推・判別")
        for q in [x for x in qs if x["lv"] == "C"]:
            question(doc, q, teacher, variant)
    if not teacher:
        closing(doc)
    tag = "基礎版" if basic else "挑戰版"
    name = f"樂學教室_期末考_{tag}_{'教師答案卷' if teacher else '學生卷'}.docx"
    doc.save(str(OUT / name))
    print("saved", name)

for v in ("basic", "challenge"):
    build(v, False)
    build(v, True)
